from tkinter import *
from workers import Worker
import time
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import os


def make_plan():
    """Создание информации о планах работников"""
    start = time.time()
    a = Input.get(1.0, 'end')
    parts = a.split('\n')
    i = 0
    while i < len(parts):
        if not parts[i]:
            parts.pop(i)
        else:
            i += 1
    blocks = make_groups(parts)
    flag = 0
    s = []
    for i in blocks:
        name = i[0].split(',')
        w = Worker()
        w.name = name[0]
        s.append(w)
        flag = 0
        for j in i:
            if j.find('https://') != -1:
                s[len(s) - 1].week_make(j, flag)
            elif j.find('недел') != -1 \
                    or j.find('след') != -1:
                flag += 1
            else:
                continue
    kol = 1
    whole = s.__len__()
    for i in s:
        i.make_tasks()
        print(f'Обрабатываем пользователя: {i.name} ({kol}/{whole})')
        kol = kol + 1
    out(s)
    end = time.time()
    print("Время работы: ", end - start)


def make_groups(s: list):
    """Функция разбивающая общий поток на блоки каждого работника"""

    blocks = [[]]
    i = 0
    k = 0
    for i in range(s.__len__()):
        if s[i].find('[') == -1 or i == 0:
            blocks[k].append(s[i])
        else:
            k += 1
            blocks.append([])
            blocks[k].append(s[i])
    return blocks


def sort_by_name(s):
    """Принцип сортировки: по имени"""
    return s.name


def add_empty_person(s):
    """Добавление работников в отпуске (без планов)"""


def out(s):
    """Создание файла и запись в него планов работников"""
    all_person = ["Арнольд Роман",
                  "Бородуля Вадим",
                  "Глазырин Юра",
                  "Городжий Дмитрий",
                  "Ефимов Игорь",
                  "Загоскин Степан",
                  "Захарова Евгения",
                  "Кобяков Александр",
                  "Марков Константин",
                  "Падалко Антон",
                  "Поздникин Сергей",
                  "Пузанов Алексей",
                  "Романов Алексей",
                  "Рубинштейн Борис",
                  "Рюмин Артём",
                  "Чубарова Наташа",
                  "Герасимов Андрей"]
    cwd = os.getcwd()
    doc = Document()
    i = 0
    templ = 'КИТ'
    for k in s:
        for j in s:
            if s.index(k) != s.index(j):
                if k.name == j.name:
                    s.remove(k)
                else:
                    continue
    for j in s:
        cont = j.name.split(' ')
        if templ in cont:
            cont.remove(templ)
        cont.reverse()
        j.name = ' '.join(cont)
    s.sort(key=sort_by_name)
    d = doc.add_paragraph()
    d.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    f = d.add_run('Отдел разработки ПО')
    f.font.size = Pt(18)
    p = doc.add_paragraph('Отчет по работе за неделю: <-- -->.План по работе на неделю: <-- -->.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = doc.add_table(rows=len(all_person) + 1, cols=4, style='Table Grid')
    table.allow_autofit = True
    cells = table.rows[i].cells
    cells[0].text = '№'
    cells[1].text = 'Сотрудники'
    cells[2].text = 'Эта неделя'
    cells[3].text = 'Следующая неделя'

    # for j in s:
    #    i += 1
    #    cells = table.rows[i].cells
    #    cells[0].text = i.__str__()
    #    cells[1].text = j.name
    #    cells[2].text = '- ' + '\n- '.join(j.this_tasks)
    #    cells[3].text = '- ' + '\n- '.join(j.next_tasks)
    for j in all_person:
        i += 1
        cells = table.rows[i].cells
        cells[0].text = i.__str__()
        cells[1].text = j
        cells[2].text = ""
        cells[3].text = ""
        for m in s:
            if j == m.name:
                cells[2].text = '- ' + '\n- '.join(m.this_tasks)
                cells[3].text = '- ' + '\n- '.join(m.next_tasks)
                break
            else:
                continue
    doc.save(cwd + '\Отчет по работе отдела разработки ПО.docx')
    Output.insert(1.0, "Формирование отчета закончено\nИнформацию об ошибках смотрите в консоли\n\n"
                       "Для материалов, которые вызвали сбои в работе или получили неправильный выходной файл "
                       "используйте папку \"Для материалов\" (формат.txt)")


def clear_text():
    """Функция очистки окна для ввода"""
    Input.delete(1.0, 'end')
    Output.delete(1.0, 'end')


def _onKeyRelease(event):
    """Блок с назначением горячих клавиш"""
    ctrl = (event.state & 0x4) != 0
    if event.keycode == 88 and ctrl and event.keysym.lower() != "x":
        event.widget.event_generate("<<Cut>>")
    if event.keycode == 86 and ctrl and event.keysym.lower() != "v":
        event.widget.event_generate("<<Paste>>")
    if event.keycode == 67 and ctrl and event.keysym.lower() != "c":
        event.widget.event_generate("<<Copy>>")
    if event.keycode == 65 and ctrl and event.keysym.lower() != "a":
        event.widget.event_generate("<<SelectAll>>")


if __name__ == '__main__':
    Window = Tk()
    Window['bg'] = "#555"
    Window.title("СФО")
    Window.geometry('1080x720')
    Window.bind("<Control-KeyPress>", _onKeyRelease)
    Input = Text(width=50, height=35, background="#fafafa")
    Output = Text(width=50, height=35, background="#fafafa")
    Make = Button(text="Создать отчет", background="#fafafa", command=make_plan, )
    Clear = Button(text="Очистить", background="#fafafa", command=clear_text, )
    Input.place(x=100, y=20)
    Output.place(x=600, y=20)
    Make.place(x=300, y=600, height=40, width=500)
    Clear.place(x=300, y=650, height=40, width=500)

    Window.mainloop()
